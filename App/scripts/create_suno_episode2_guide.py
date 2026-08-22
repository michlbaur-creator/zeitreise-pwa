from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from create_suno_guide import (
    BLACK,
    BLUE,
    DARK_BLUE,
    GOLD,
    LIGHT_BLUE,
    MID_GRAY,
    NAVY,
    WHITE,
    add_callout,
    add_page_number,
    add_prompt_box,
    set_repeat_table_header,
    set_run,
    set_cell_margins,
    set_table_borders,
    set_table_geometry,
    shade,
)


OUTPUT = (
    Path(__file__).resolve().parents[2]
    / "Dokumente"
    / "03_Arbeitsstand"
    / "Suno-Pro_Episode2_Geraeuscheliste.docx"
)

LOOP_SUFFIX = (
    "natural history documentary field recording, seamless loop, 25 seconds, "
    "subtle and calm, steady volume, no music, no melody, no human voices, "
    "no narration, no dramatic cinematic effects, no modern sounds"
)

ONE_SHOT_SUFFIX = (
    "single isolated sound event, clean background, no repetition, no music, "
    "no voices, no cinematic score, no modern sounds"
)

AMBIENCES = [
    {
        "filename": "atmo_ep2_waldkronen_loop.wav",
        "scenes": "Szenen 1–2",
        "label": "Waldkronen der frühen Primaten",
        "prompt": (
            "warm humid African forest canopy, gentle irregular leaf rustle, "
            "soft branch movement, sparse distant insects and only a few small birds, "
            "spacious and unobtrusive, no identifiable primate calls, " + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_waldsavanne_loop.wav",
        "scenes": "Szenen 3–4",
        "label": "Afrikanisches Wald-Savannen-Mosaik",
        "prompt": (
            "African woodland and savanna mosaic, gentle breeze through scattered trees "
            "and low grass, sparse insects, very occasional distant birds, dry but not harsh, "
            "no large animal calls, " + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_grasland_loop.wav",
        "scenes": "Szenen 5–8",
        "label": "Offenes afrikanisches Grasland",
        "prompt": (
            "open African grassland, soft dry wind moving through low grass, faint distant "
            "insects and very sparse birds, broad open space, no footsteps, no animal calls, "
            + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_feuer_lager_loop.wav",
        "scenes": "Szene 9",
        "label": "Kontrolliertes Feuer",
        "prompt": (
            "small controlled prehistoric campfire outdoors, gentle irregular crackling, "
            "occasional soft ember shift, faint evening air, intimate but restrained, "
            "no people, no cooking noises, " + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_kaltes_lager_loop.wav",
        "scenes": "Szenen 10–11",
        "label": "Kühles offenes Lager",
        "prompt": (
            "cold open Eurasian landscape near a quiet camp, low steady wind, subtle fabric "
            "and grass movement, very faint distant fire crackle, no snowstorm, no people, "
            + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_hoehle_loop.wav",
        "scenes": "Szene 12",
        "label": "Ruhige Felshöhle",
        "prompt": (
            "quiet limestone cave interior, subtle natural air movement, rare soft water drop, "
            "short realistic cave reflection, deep stillness, no drone, no animals, "
            + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_wasserstelle_loop.wav",
        "scenes": "Szene 13",
        "label": "Afrikanische Wasserstelle",
        "prompt": (
            "calm African waterhole, gentle water movement at a stony shore, soft breeze in "
            "dry vegetation, sparse insects and only a few distant birds, no splashing animals, "
            + LOOP_SUFFIX
        ),
    },
    {
        "filename": "atmo_ep2_finale_loop.wav",
        "scenes": "Szene 14",
        "label": "Offene Landschaft im Finale",
        "prompt": (
            "wide peaceful African landscape in warm late daylight, soft open wind, a few "
            "distant small birds, faint vegetation movement, reflective and spacious, "
            "no triumphant mood, " + LOOP_SUFFIX
        ),
    },
]

ONE_SHOTS = [
    {
        "filename": "geraeusch_ep2_ascheschritte.wav",
        "scenes": "Szene 5",
        "label": "Schritte in Vulkanasche",
        "cue": "Beim ersten deutlichen Hinweis auf die Fußspuren einblenden.",
        "prompt": (
            "short sequence of three soft bare human footsteps pressing into fine damp volcanic "
            "ash, close natural foley, gentle weight, no running, " + ONE_SHOT_SUFFIX
        ),
    },
    {
        "filename": "geraeusch_ep2_steinschlag.wav",
        "scenes": "Szene 6",
        "label": "Ein einzelner Steinschlag",
        "cue": "Genau auf die gesprochenen Worte „Ein Schlag“ setzen.",
        "prompt": (
            "one heavy stone striking another stone once, dry sharp impact, a few small stone "
            "fragments, short natural tail, no metallic ring, no echo, " + ONE_SHOT_SUFFIX
        ),
    },
]

MUSIC = [
    {
        "filename": "musik_ep2_anfang.wav",
        "scenes": "Szene 1 · etwa 10–14 Sekunden",
        "label": "Kurzer Beginn",
        "prompt": (
            "short restrained instrumental opening for a natural history documentary about "
            "human evolution, warm low strings, soft wooden and stone textures, a gentle sense "
            "of curiosity, very slow, no vocals, no choir, no drums, no trailer climax, "
            "quiet ending, approximately 14 seconds"
        ),
    },
    {
        "filename": "musik_ep2_mitteluebergang.wav",
        "scenes": "Szene 8 · etwa 16–20 Sekunden",
        "label": "Großer Mittelübergang",
        "prompt": (
            "restrained instrumental transition for a natural history documentary, suggesting "
            "a long human journey across landscapes, slow airy strings, low warm woodwind "
            "textures, one gentle rising arc, no vocals, no choir, no drums, no heroic trailer "
            "sound, soft open ending, approximately 18 seconds"
        ),
    },
    {
        "filename": "musik_ep2_finale.wav",
        "scenes": "Szene 14 · letzte 22–28 Sekunden",
        "label": "Ruhiges Finale",
        "prompt": (
            "quiet reflective instrumental finale for a natural history documentary about one "
            "shared humanity, soft sustained strings, muted piano and a faint airy woodwind "
            "texture, warm and thoughtful, no vocals, no choir, no percussion, not triumphant, "
            "gentle resolved ending, approximately 26 seconds"
        ),
    },
]

SCENE_MAP = [
    ("01", "Der nächste Zeitsprung", "Waldkronen", "atmo_ep2_waldkronen_loop.wav + musik_ep2_anfang.wav"),
    ("02", "Die Welt der Primaten", "Waldkronen", "atmo_ep2_waldkronen_loop.wav"),
    ("03", "Getrennte Wege", "Wald-Savanne", "atmo_ep2_waldsavanne_loop.wav"),
    ("04", "Der rätselhafte Gang", "Wald-Savanne", "atmo_ep2_waldsavanne_loop.wav"),
    ("05", "Lucy und die Spuren", "Grasland + Schritte", "atmo_ep2_grasland_loop.wav + geraeusch_ep2_ascheschritte.wav"),
    ("06", "Der Stein", "Grasland + Steinschlag", "atmo_ep2_grasland_loop.wav + geraeusch_ep2_steinschlag.wav"),
    ("07", "Homo kommt in Bewegung", "Grasland", "atmo_ep2_grasland_loop.wav"),
    ("08", "Die erste große Reise", "Grasland + Musik", "atmo_ep2_grasland_loop.wav + musik_ep2_mitteluebergang.wav"),
    ("09", "Feuer verändert den Alltag", "Feuer", "atmo_ep2_feuer_lager_loop.wav"),
    ("10", "Viele Menschenformen", "Kühles Lager", "atmo_ep2_kaltes_lager_loop.wav"),
    ("11", "Die Neandertaler", "Kühles Lager", "atmo_ep2_kaltes_lager_loop.wav"),
    ("12", "Denisova", "Höhle", "atmo_ep2_hoehle_loop.wav"),
    ("13", "Homo sapiens trifft Verwandte", "Wasserstelle", "atmo_ep2_wasserstelle_loop.wav"),
    ("14", "Eine Menschheit", "Finale + Musik", "atmo_ep2_finale_loop.wav + musik_ep2_finale.wav"),
]


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 14, MID_GRAY, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name != "Subtitle" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    run = header.add_run("ZEITREISE  |  EPISODE 2  |  PRODUKTIONSHILFE")
    set_run(run, size=8.5, color=MID_GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])
    return doc


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(15)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("VOLLSTÄNDIGE PRODUKTIONSLISTE"), size=10, color=GOLD, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Geräusche mit Suno Pro erstellen")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Episode 2 · Die Entwicklung des Menschen")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    set_run(p.add_run("8 Atmosphären · 2 Einzelgeräusche · 3 Musikmomente\n"), size=11.5, color=NAVY, bold=True)
    set_run(p.add_run("Stand: August 2026"), size=10, color=MID_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    set_run(p.add_run("Die Stimme führt. Der Klang trägt nur die Landschaft."), size=11, color=GOLD, italic=True)
    doc.add_page_break()


def add_scene_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for index, text in enumerate(("Szene", "Titel", "Klangidee", "Dateien")):
        cell = table.rows[0].cells[index]
        set_cell_margins(cell, top=35, start=80, bottom=35, end=80)
        shade(cell, NAVY)
        set_run(cell.paragraphs[0].add_run(text), size=9.1, color=WHITE, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(SCENE_MAP):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                shade(cell, "F7F9FB")
        for index, value in enumerate(row):
            set_cell_margins(cells[index], top=24, start=80, bottom=24, end=80)
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run(p.add_run(value), size=8.2 if index == 3 else 8.5, bold=index == 0, font="Aptos Mono" if index == 3 else "Aptos")
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, [720, 2300, 1900, 4440])
    set_table_borders(table)


def add_asset(doc, asset, number, music=False):
    doc.add_heading(f"{number}. {asset['label']}", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Dateiname: "), size=9.5, color=DARK_BLUE, bold=True)
    set_run(p.add_run(asset["filename"]), size=9.5, font="Aptos Mono")
    set_run(p.add_run("   Verwendung: "), size=9.5, color=DARK_BLUE, bold=True)
    set_run(p.add_run(asset["scenes"]), size=9.5)
    if asset.get("cue"):
        c = doc.add_paragraph()
        set_run(c.add_run("Einsatz: "), size=9.5, color=DARK_BLUE, bold=True)
        set_run(c.add_run(asset["cue"]), size=9.5)
    label = "Suno · Create · Custom · Instrumental" if music else "Suno Sounds · Prompt"
    add_prompt_box(doc, label, asset["prompt"])


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    add_cover(doc)

    doc.add_heading("1. Klangkonzept", level=1)
    doc.add_paragraph(
        "Episode 2 benötigt keine lückenlose Filmmusik. Michas Sprecheraufnahme bleibt der klare Mittelpunkt. "
        "Darunter liegen sehr leise Naturkulissen; nur zwei kurze Ereignisse werden punktgenau betont. Musik erscheint "
        "ausschließlich am Anfang, beim großen Mittelübergang und im Finale."
    )
    add_callout(doc, "Leitregel:", "Wenn ein Geräusch bewusst auffällt, ist es wahrscheinlich bereits zu laut. Zuerst immer die ruhigste und natürlichste Suno-Variante auswählen.")

    doc.add_heading("2. Schnellstart in Suno", level=1)
    steps = [
        "Suno öffnen und zu Create wechseln.",
        "Für Atmosphären Loop und für Einzelereignisse One Shot im Bereich Sounds wählen.",
        "Für die drei Musikstücke Custom und Instrumental verwenden.",
        "Den vollständigen englischen Prompt aus dieser Liste kopieren.",
        "Zwei oder mehr Varianten erzeugen und mit Kopfhörern vergleichen.",
        "Die beste Fassung als WAV laden und exakt mit dem angegebenen Dateinamen speichern.",
    ]
    for step in steps:
        paragraph = doc.add_paragraph(step, style="List Number")
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05

    mapping_heading = doc.add_heading("3. Vollständige Zuordnung zu allen 14 Szenen", level=1)
    mapping_heading.paragraph_format.space_before = Pt(10)
    add_scene_table(doc)
    doc.add_page_break()

    doc.add_heading("4. Atmosphären und Endlosschleifen", level=1)
    doc.add_paragraph("Diese acht Dateien decken alle 14 Szenen ab. Wiederverwendung ist beabsichtigt und sorgt für einen zusammenhängenden Klang.")
    for index, asset in enumerate(AMBIENCES, 1):
        if index in (5,):
            doc.add_page_break()
        add_asset(doc, asset, index)

    doc.add_page_break()
    doc.add_heading("5. Einzelgeräusche", level=1)
    doc.add_paragraph("Nur zwei Ereignisse brauchen eine eigene, punktgenaue Betonung. Weitere künstliche Effekte würden die Episode eher überladen.")
    for index, asset in enumerate(ONE_SHOTS, 1):
        add_asset(doc, asset, index)

    doc.add_heading("6. Die drei Musikmomente", level=1)
    doc.add_paragraph("Diese Prompts im normalen Suno-Bereich mit eingeschaltetem Instrumental-Modus verwenden. Benötigt wird jeweils nur der passende kurze Ausschnitt.")
    for index, asset in enumerate(MUSIC, 1):
        add_asset(doc, asset, index, music=True)

    doc.add_heading("7. Auswahlregeln", level=1)
    rules = [
        "Keine Stimmen, Rufe, Gesänge oder verständlichen menschlichen Laute.",
        "Keine rhythmischen Schläge unter dem Sprecher.",
        "Tierstimmen nur entfernt und sehr sparsam; keine modernen Stadt- oder Verkehrsgeräusche.",
        "Loops mehrfach hintereinander anhören: Der Übergang darf nicht auffallen.",
        "Bei Szene 6 genau einen Steinschlag verwenden.",
        "Die Musik darf weder heroisch noch traurig noch wie ein Filmtrailer wirken.",
        "Atmosphäre und Musik deutlich leiser als Michas Stimme halten.",
    ]
    for rule in rules:
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("8. Ablage und Übergabe", level=1)
    doc.add_paragraph("Die fertigen WAV-Dateien bitte gemeinsam in folgenden Ordner legen:")
    add_prompt_box(doc, "Ordnername", "Austausch/ZeitreiseTon2-Sound")
    for item in (
        "Dateinamen aus dieser Liste unverändert übernehmen.",
        "Keine Dateien selbst schneiden oder normalisieren; das erfolgt beim Einbau in die App.",
        "Zu jeder gewählten Datei nach Möglichkeit Suno-Link und Erstellungsdatum notieren.",
        "Nach dem Einbau Episode 2 einmal mit Kopfhörern und einmal über den iPhone-Lautsprecher prüfen.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    add_callout(doc, "Nächster Schritt:", "Sobald die Dateien im Austausch-Ordner liegen, werden Lautstärke, Überblendungen und die beiden exakten Einsatzpunkte in der App eingerichtet.", fill=LIGHT_BLUE)

    doc.add_heading("9. Offizielle Suno-Hinweise", level=1)
    doc.add_paragraph("Suno Sounds ist derzeit als experimentelle Funktion für Geräusche, Samples und Atmosphären beschrieben. Für verwendete Dateien aus einem bezahlten Tarif gelten Sunos aktuelle Nutzungsbedingungen; die rechtliche Schutzfähigkeit kann davon getrennt sein.")
    for source in (
        "Suno Sounds: https://help.suno.com/en/articles/10625537",
        "Musik im Instrumental-Modus: https://help.suno.com/en/articles/2462273",
        "Nutzungsrechte: https://help.suno.com/en/articles/2746945",
    ):
        doc.add_paragraph(source, style="List Bullet")

    props = doc.core_properties
    props.title = "Suno Pro – Geräuscheliste für Zeitreise Episode 2"
    props.subject = "Vollständige Produktionsliste für Atmosphären, Einzelgeräusche und Musikmomente"
    props.author = "Zeitreise – Die Geschichte des Lebens"
    props.keywords = "Suno, Episode 2, Geräusche, Musik, Sounddesign, Zeitreise"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
