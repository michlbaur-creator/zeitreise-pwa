from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("Dokumente/03_Arbeitsstand/Zeitreise_Episode2_Sprechertexte_Humorfassung_V0.2.docx")

SCENES = [
    (1, "Der nächste Zeitsprung", "Vor etwa 60 bis 55 Millionen Jahren", "ca. 38 Sekunden",
     "Unsere erste Reise endete bei einem kleinen Säugetier. Genau dort geht es weiter. Wälder breiten sich aus, und in den Baumkronen werden geschickte Hände, gutes räumliches Sehen und ein sicherer Gleichgewichtssinn nützlich. Wer hoch oben daneben greift, bekommt schließlich keine zweite Probe. Dieses Tier ist trotzdem nicht einfach ‚der erste Mensch‘. So bequem macht es uns die Evolution nicht. Es gehört zu einer verzweigten Verwandtschaft, aus der später Lemuren, Affen, Menschenaffen und auch wir hervorgehen. Willkommen in Episode 2 – einer Geschichte mit vielen Ästen, einigen Sackgassen und ganz ohne Aufzug nach oben."),
    (2, "Ein Leben in den Bäumen", "Vor etwa 55 bis 35 Millionen Jahren", "ca. 36 Sekunden",
     "In den Bäumen zählt jeder Griff. Wer eine Entfernung falsch einschätzt, landet nicht neben dem Weg – unter dem Weg geht es weit nach unten. Viele frühe Primaten besitzen deshalb bewegliche Hände und Füße, empfindliche Fingerspitzen und nach vorn gerichtete Augen, die Entfernungen gut erfassen. Eine Art eingebautes Abstandsmessgerät, nur ohne Batterieanzeige. Doch es gibt keinen perfekten Bauplan: Manche Arten sind klein und nachtaktiv, andere werden größer und probieren neue Nahrung aus. Evolution arbeitet nicht auf ein Ziel hin. Sie behält, was unter den jeweiligen Bedingungen funktioniert."),
    (3, "Die Welt der Menschenaffen", "Vor etwa 23 bis 10 Millionen Jahren", "ca. 40 Sekunden",
     "Jetzt wird der Stammbaum richtig buschig. Im Miozän leben viele verschiedene Menschenaffen in Afrika, Asien und Europa. Manche klettern geschickt, manche bewegen sich häufiger am Boden, und viele verschwinden wieder, ohne direkte Nachkommen zu hinterlassen. Keine dieser Formen trägt ein Schild mit der Aufschrift: ‚Hier entlang zum Menschen.‘ Unsere Linie ist nur einer von vielen Zweigen. Und wie bei einem echten Baum gilt: Ein Ast wird nicht besser, nur weil er später weiterwächst. Evolution verteilt keine Pokale – sie produziert Vielfalt. Wer daraus eine ordentliche Leiter machen möchte, braucht vor allem eines: sehr viel Fantasie."),
    (4, "Afrika im Wandel", "Vor etwa 10 bis 7 Millionen Jahren", "ca. 38 Sekunden",
     "Afrika wird nicht plötzlich zu einer einzigen endlosen Savanne. Die Natur hat schließlich keinen Landschaftsgärtner bestellt. Wälder, Buschland, Flussufer und offene Flächen verschieben immer wieder ihre Grenzen. Wer hier lebt, muss mit Veränderungen zurechtkommen: beim Fortbewegen, bei der Nahrungssuche und vielleicht auch beim Zusammenarbeiten. Dieses wechselnde Mosaik könnte neue Fähigkeiten begünstigt haben. Einen einzelnen Schalter mit der Aufschrift ‚Menschwerdung‘ finden Forschende jedoch nicht. Klima, Lebensraum, Körperbau und Verhalten beeinflussen sich über sehr lange Zeit gegenseitig. Es ist ein kompliziertes Zusammenspiel – eher Mischpult als Lichtschalter."),
    (5, "Getrennte Wege", "Vor ungefähr 9 bis 6 Millionen Jahren", "ca. 42 Sekunden",
     "Ein wichtiger Satz für unsere Reise: Menschen stammen nicht von den heute lebenden Schimpansen ab. Wir teilen mit ihnen einen älteren gemeinsamen Vorfahren. Wie dieser genau aussah, wissen wir nicht – Familienfotos waren damals noch erstaunlich selten. Irgendwann vor ungefähr neun bis sechs Millionen Jahren trennen sich Populationen, und ihre Nachkommen entwickeln sich auf verschiedenen Wegen weiter. Auch Schimpansen und Bonobos haben seitdem eine lange eigene Geschichte. Die Gabelung führt also nicht von ‚primitiv‘ zu ‚fortschrittlich‘. Sie ist der Beginn mehrerer gleichzeitiger Geschichten. Wir sind nicht das Ziel dieses Stammbaums, sondern einer seiner jüngeren Zweige."),
    (6, "Auf zwei Beinen?", "Vor etwa 7 bis 6 Millionen Jahren", "ca. 42 Sekunden",
     "War das schon ein aufrechter Gang? Sahelanthropus lebte vor ungefähr sieben Millionen Jahren, Orrorin etwas später. Untersuchungen von Knochen stärken die Annahme, dass sie sich zumindest zeitweise auf zwei Beinen bewegten. Elegant wie bei einem heutigen Spaziergang sah das vermutlich noch nicht aus. Ihre Arme blieben jedenfalls gut zum Klettern geeignet – sicher ist sicher. Doch die Fossilien sind unvollständig, und ihre genaue Stellung im Stammbaum bleibt umstritten. Das ist kein Fehler der Wissenschaft. Im Gegenteil: Wissenschaft zeigt nicht nur Antworten, sondern auch, wie belastbar sie sind. Manchmal lautet die ehrlichste Antwort eben: ziemlich wahrscheinlich, aber noch nicht endgültig."),
    (7, "Ardi – zwischen Baum und Boden", "Vor etwa 4,4 Millionen Jahren", "ca. 42 Sekunden",
     "Ardi passt in keine einfache Schublade – und würde vermutlich auch sofort wieder herausklettern. Sie lebt vor rund 4,4 Millionen Jahren in einer waldreichen Landschaft. Ihr Becken zeigt Merkmale, die zu aufrechtem Gehen passen könnten. Gleichzeitig besitzt ihr Fuß einen greiffähigen großen Zeh, äußerst praktisch zwischen Ästen. Ihr Körper verbindet also verschiedene Möglichkeiten: Gehen am Boden und Klettern im Baum. Zweibeinigkeit entsteht nicht als fertige Erfindung an einem einzigen Tag. Sie entwickelt sich schrittweise, während ältere Fähigkeiten weiter nützlich bleiben. Evolution baut selten alles neu. Sie arbeitet lieber mit dem, was schon vorhanden ist."),
    (8, "Spuren in der Asche", "Vor rund 3,6 Millionen Jahren", "ca. 38 Sekunden",
     "Knochen zeigen, wie ein Körper gebaut war. Diese Spuren zeigen sogar, was jemand getan hat. Vor rund 3,6 Millionen Jahren laufen frühe Homininen durch feuchte Vulkanasche in Laetoli – vermutlich ohne zu ahnen, dass ihre Fußabdrücke einmal weltberühmt werden. Die großen Zehen liegen in einer Reihe mit den übrigen, und der Fuß rollt von der Ferse nach vorn ab. Wahrscheinlich stammen die Spuren von Australopithecus afarensis. Wer genau hier ging und wohin, wissen wir nicht. Doch für einen Moment wird aus einem fernen Fossil eine Bewegung: Schritt für Schritt."),
    (9, "Lucy – zu Hause in zwei Welten", "Vor etwa 3,2 Millionen Jahren", "ca. 44 Sekunden",
     "Lucy ist nur etwa einen Meter groß, aber ihr Skelett erzählt eine ziemlich große Geschichte. Ihr Becken und ihre Beine zeigen einen Körper, der regelmäßig aufrecht geht. Lange Arme und gebogene Finger verraten zugleich, dass Klettern weiterhin wichtig ist. Warum auch eine bewährte Fähigkeit aufgeben, nur weil man jetzt häufiger zu Fuß unterwegs ist? Ihr Gehirn ist deutlich kleiner als unseres. Erst aufrecht gehen und viel später ein großes Gehirn entwickeln – diese Reihenfolge überrascht viele. Lucy ist dabei nicht ‚halb Affe, halb Mensch‘. Sie ist vollständig sie selbst: Australopithecus afarensis, angepasst an Boden und Bäume – zwei Welten zum Preis von einer."),
    (10, "Der Stein wird zum Werkzeug", "Früheste bekannte Funde vor etwa 3,3 Millionen Jahren", "ca. 48 Sekunden",
     "Ein Schlag – und der Stein verändert sich. In Lomekwi wurden Werkzeuge gefunden, die rund 3,3 Millionen Jahre alt sind. Damit sind sie älter als die frühesten sicher bekannten Fossilien der Gattung Homo. Wer sie hergestellt hat, wissen wir nicht. Vielleicht eine Australopithecus-Art, vielleicht eine andere Homininenform. Der Name des Erfinders ist jedenfalls nicht eingeritzt. Sicher ist: Werkzeugtechnik fällt nicht plötzlich fertig vom Himmel. Jemand wählt Material aus, hält einen Stein fest und verändert ihn mit wiederholten Schlägen. Das Ergebnis wirkt für uns grob. Doch dahinter stecken Planung, Kraft und Übung – vermutlich auch einige Schläge, die völlig daneben gingen."),
    (11, "Die Gattung Homo", "Seit ungefähr 2,8 Millionen Jahren", "ca. 42 Sekunden",
     "Mit der Gattung Homo beginnt kein sauber abgegrenztes neues Kapitel. Die Evolution blättert nicht um und schreibt oben ordentlich ‚Jetzt kommt der Mensch‘. Der älteste bekannte Kiefer, der zu Homo gestellt wird, ist etwa 2,8 Millionen Jahre alt. Später begegnen uns Namen wie Homo habilis. Doch Zähne, Hände, Gehirn und Körper verändern sich nicht alle im selben Tempo. Gleichzeitig leben weiterhin Australopithecus- und Paranthropus-Arten. Statt einer feierlichen Wachablösung sehen wir mehrere Zweige nebeneinander – manche kurz, manche erstaunlich langlebig. Die Schubladen stammen von uns. Die Natur hält sich nur leider nicht immer an unsere Beschriftungen."),
    (12, "Homo erectus – ein Körper für lange Wege", "Ab etwa 1,9 Millionen Jahren", "ca. 44 Sekunden",
     "Homo erectus wirkt im Körperbau schon deutlich vertrauter: längere Beine, kürzere Arme und ein Körper, der weite Strecken besser bewältigt. Eine gute Voraussetzung, wenn der nächste brauchbare Wasserplatz nicht gleich hinter dem ersten Busch liegt. Diese Menschenform erscheint vor ungefähr 1,9 Millionen Jahren und bleibt in verschiedenen Regionen sehr lange bestehen. Ihr Gehirn ist im Durchschnitt größer als bei früheren Homininen, aber noch nicht wie unseres. Entscheidend ist ohnehin das Zusammenspiel: Körper, Nahrung, Werkzeuge, Lernen und soziale Gruppen eröffnen gemeinsam neue Möglichkeiten. Kein einzelnes Merkmal macht plötzlich den Menschen – auch ein langes Bein reist selten allein."),
    (13, "Die erste große Reise", "Vor etwa 1,85 bis 1,7 Millionen Jahren", "ca. 46 Sekunden",
     "Vor ungefähr 1,8 Millionen Jahren finden wir frühe Menschen außerhalb Afrikas. In Dmanisi im heutigen Georgien leben kleine Gruppen zwischen Grasland, Wald und Bergen. Später reichen Spuren bis weit nach Ostasien. Wahrscheinlich ist das keine einzige große Auswanderung mit gepacktem Reiseproviant und festem Ziel. Generation für Generation folgen Menschen Nahrung, Wasser und günstigen Lebensräumen. Manche Populationen kehren um, andere verschwinden, wieder andere breiten sich weiter aus. Auf einer Karte sieht das wie eine große Reise aus. Für die Beteiligten war es vermutlich immer nur der nächste erreichbare Weg – und gelegentlich die Frage: Wo gibt es heute etwas zu essen?"),
    (14, "Feuer verändert den Alltag", "Sicher kontrollierte Feuerstellen spätestens vor etwa 790.000 Jahren", "ca. 45 Sekunden",
     "Feuer ist mehr als Wärme. Es macht manche Nahrung leichter verdaulich, spendet Licht, kann Tiere fernhalten und schafft einen Treffpunkt. Vermutlich war es auch der erste Ort, an dem alle gleichzeitig behaupteten, jemand anderes solle Holz nachlegen. Sicher nachgewiesene Feuerstellen sind mindestens etwa 790.000 Jahre alt. Es gibt ältere Hinweise, doch sie sind schwieriger zu deuten. Ein verbrannter Fleck ist schließlich noch keine Küche – vielleicht hat dort einfach ein Busch gebrannt. Entscheidend ist, ob Menschen Feuer wiederholt und kontrolliert nutzen. Erst dann wird aus einer Flamme ein Werkzeug, das Nahrung, Schutz und gemeinsames Leben nachhaltig verändert."),
    (15, "Viele Arten von Menschen", "Vor etwa 700.000 bis 300.000 Jahren", "ca. 48 Sekunden",
     "Je näher wir unserer Gegenwart kommen, desto voller wird die Welt der Menschenformen. Populationen leben in Afrika, Europa und Asien, trennen sich, begegnen sich später wieder und verändern sich weiter. Ältere Bücher ordneten viele Funde gern in eine saubere Reihe – das sieht übersichtlich aus, ist aber wahrscheinlich zu ordentlich. Heute ist das Bild vorsichtiger. Manche Fossilien lassen sich nur schwer einer Art zuordnen, und selbst Bezeichnungen wie Homo heidelbergensis werden unterschiedlich verwendet. Die Menschenformen kannten unsere Namensschilder schließlich nicht. Sicher ist vor allem die Verzweigung: Viele Populationen lebten gleichzeitig, und nicht jede wissenschaftliche Schublade schließt so sauber, wie wir es gern hätten."),
    (16, "Die Neandertaler", "Vor etwa 400.000 bis 40.000 Jahren", "ca. 50 Sekunden",
     "Neandertaler sind keine unbeholfenen Vorstufen von uns – und ganz sicher keine Statisten, die nur auf unseren Auftritt warten. Sie leben Hunderttausende Jahre in Europa und Westasien, stellen anspruchsvolle Werkzeuge her, nutzen Feuer und kümmern sich offenbar auch um verletzte oder kranke Gruppenmitglieder. Ihr kräftiger Körper passt gut zu kalten Bedingungen. Wer im eiszeitlichen Europa lebt, kann schließlich nicht einfach die Heizung höher drehen. Gleichzeitig verändern sich ihre Kulturen über Zeit und Raum. Neandertaler sind ein eigener, erfolgreicher Menschenzweig. Und völlig getrennt von uns waren sie ebenfalls nicht. Davon tragen viele Menschen heute noch winzige Spuren im eigenen Erbgut."),
    (17, "Die Denisova-Menschen", "Vor mehreren hunderttausend bis vor mindestens etwa 50.000 Jahren", "ca. 55 Sekunden",
     "Manchmal passt eine ganze Menschenlinie in eine winzige Schachtel. Denisova-Menschen wurden zuerst durch DNA aus einem kleinen Fingerknochen erkannt – weniger Fundmaterial, als man für ein ordentliches Puzzle erwarten würde. Lange kannten wir fast nur Zähne und Bruchstücke. Seit 2025 liefern alte Proteine und Erbgut aus Zahnstein einen größeren Hinweis: Auch der mehr als 146.000 Jahre alte Harbin-Schädel gehört sehr wahrscheinlich zu dieser Linie. Das Bild wird klarer, bleibt aber unvollständig. Genetische Spuren zeigen, dass Denisova-Populationen weit in Asien lebten und sich mit Neandertalern und Homo sapiens vermischten. Ein Mädchen mit Neandertaler-Mutter und Denisova-Vater macht solche Begegnungen besonders greifbar. Die DNA erzählt hier mehr als die Knochenkiste."),
    (18, "Homo sapiens entsteht", "Seit mindestens etwa 300.000 Jahren", "ca. 50 Sekunden",
     "Unsere eigene Art erscheint nicht plötzlich an einem einzigen Geburtsort – es gab weder eine Wiege mit Namensschild noch einen feierlichen ersten Geburtstag. Fossilien aus Jebel Irhoud in Marokko sind rund 300.000 Jahre alt. Weitere frühe Funde stammen aus anderen Teilen Afrikas. Manche Merkmale wirken schon modern, andere noch älter. Viele Forschende sprechen deshalb von einer panafrikanischen Entstehung: Populationen in verschiedenen Regionen bleiben über lange Zeit miteinander verbunden und tragen gemeinsam zur Entwicklung von Homo sapiens bei. Auch wir sind also das Ergebnis eines Netzes, nicht eines einzigen Punktes. Das ist komplizierter als eine Ursprungsgeschichte mit Stecknadel auf der Karte – aber sehr viel spannender."),
    (19, "Begegnungen und geteilte Geschichten", "Vor etwa 120.000 bis 40.000 Jahren", "ca. 54 Sekunden",
     "Als Homo sapiens sich weiter ausbreitet, ist die Welt nicht leer. Menschen treffen auf Neandertaler und Denisova-Populationen. Manche Gruppen konkurrieren vielleicht, manche gehen einander aus dem Weg, andere bekommen gemeinsame Kinder. Genetische Spuren zeigen uns solche Verbindungen – ob dabei jemand freundlich gegrüßt hat, verraten sie leider nicht. Das Erbgut vieler heute lebender Menschen bewahrt Teile dieser Begegnungen. Gleichzeitig entstehen in verschiedenen Regionen neue Werkzeuge, Schmuck, Bilder und weitreichende Netzwerke. Es gibt keinen einzelnen Morgen, an dem ‚moderne Kultur‘ plötzlich eingeschaltet wird. Fähigkeiten wachsen, wandern, verschwinden und werden neu kombiniert. Menschheitsgeschichte ist kein Staffellauf mit sauberer Übergabe, sondern ein Geflecht aus vielen Wegen."),
    (20, "Eine Menschheit", "Heute", "ca. 55 Sekunden",
     "Unsere Geschichte ist verzweigt, voller Seitenwege und Begegnungen. Viele Menschenformen leben lange gleichzeitig. Einige verschwinden, andere hinterlassen Spuren in unserem Erbgut. Heute gibt es nur noch eine Menschenart: Homo sapiens. Unsere sichtbare Vielfalt ist real und wertvoll, doch sie teilt uns nicht in biologische Menschenrassen. Dafür sind wir viel zu eng miteinander verwandt – selbst wenn Familienfeiern gelegentlich einen anderen Eindruck vermitteln. Letztlich stammen wir alle aus afrikanischen Populationen. Deine Hand liegt wieder auf dem Zeitfelsen. Die Reise endet hier – vorerst. Denn was Menschen später mit ihrer Welt machen, ist ein neues Kapitel. Und wie du dir denken kannst, wird auch dieses Kapitel nicht gerade kurz."),
]


NAVY = RGBColor(32, 55, 72)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 99, 104)
GOLD = RGBColor(174, 128, 61)
PALE = "F7F3E8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), str(width_dxa))
    grid.append(grid_col)
    for cell in table.columns[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        tc_w.set(qn("w:w"), str(width_dxa))
        tc_w.set(qn("w:type"), "dxa")
        margins = OxmlElement("w:tcMar")
        for side, value in (("top", 150), ("bottom", 150), ("start", 180), ("end", 180)):
            node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")
            margins.append(node)
        tc_pr.append(margins)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_run(run, size=11, color=None, bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_run(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "ZEITREISE  ·  EPISODE 2  ·  SPRECHERTEXTE"
    set_run(header.runs[0], size=8.5, color=MUTED, bold=True)
    header.paragraph_format.space_after = Pt(0)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(kicker.add_run("ZEITREISE – DIE GESCHICHTE DES LEBENS"), size=10, color=GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(title.add_run("Episode 2"), size=30, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(5)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("Die Entwicklung des Menschen"), size=19, color=BLUE)
    subtitle.paragraph_format.space_after = Pt(3)
    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(strap.add_run("Eine verzweigte Geschichte"), size=14, color=MUTED, italic=True)
    strap.paragraph_format.space_after = Pt(28)
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(version.add_run("Sprechertexte – Humorfassung V0.2"), size=11, color=GOLD, bold=True)
    version.paragraph_format.space_after = Pt(88)
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(status.add_run("Prüffassung zur redaktionellen Freigabe\nNoch nicht in die App übernommen"), size=10, color=MUTED, italic=True)
    doc.add_page_break()


def add_intro(doc):
    doc.add_heading("Hinweise zur Fassung", level=1)
    p = doc.add_paragraph()
    p.add_run("Ziel: ").bold = True
    p.add_run("Die wissenschaftlich freigegebene Geschichte bleibt unverändert, erhält aber eine wärmere, etwas humorvollere Erzählstimme.")
    points = [
        "Humor entsteht aus Vergleichen, überraschenden Bildern und kleinen trockenen Bemerkungen.",
        "Keine erfundenen Dialoge, keine Urmenschen-Karikaturen und keine Witze auf Kosten früher Menschenformen.",
        "Unsicherheiten bleiben ausdrücklich als Unsicherheiten erkennbar.",
        "Szene 19 und das Finale bleiben ruhiger; das Augenzwinkern tritt dort bewusst zurück.",
        "Zeitangaben, Szenenfolge und wissenschaftliche Kernaussagen wurden nicht verändert.",
    ]
    for text in points:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(14)
    set_run(note.add_run("Redaktioneller Status: "), size=10.5, color=GOLD, bold=True)
    set_run(note.add_run("Diese Fassung ersetzt die verbindlichen Sprechertexte erst nach ausdrücklicher Freigabe."), size=10.5, color=MUTED)
    doc.add_page_break()


def add_scene(doc, index, title, time_label, duration, speaker):
    heading = doc.add_heading(f"{index}. {title}", level=1)
    heading.paragraph_format.keep_with_next = True
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(9)
    meta.paragraph_format.keep_with_next = True
    set_run(meta.add_run(time_label), size=9.5, color=GOLD, bold=True)
    set_run(meta.add_run(f"   ·   {duration}   ·   {len(speaker.split())} Wörter"), size=9.5, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(speaker), size=11.2, color=NAVY)

    review = doc.add_paragraph()
    review.paragraph_format.space_before = Pt(7)
    review.paragraph_format.space_after = Pt(8)
    set_run(review.add_run("Prüfhinweis: "), size=8.5, color=GOLD, bold=True)
    set_run(review.add_run("Ton, Pointe und Sprechlänge gemeinsam freigeben."), size=8.5, color=MUTED, italic=True)


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_intro(doc)
    for scene in SCENES:
        add_scene(doc, *scene)
    props = doc.core_properties
    props.title = "Zeitreise – Episode 2 – Sprechertexte Humorfassung V0.2"
    props.subject = "Prüffassung der Sprechertexte für Episode 2"
    props.author = "Michael Baur / Codex"
    props.keywords = "Zeitreise, Evolution, Menschheitsentwicklung, Sprechertexte"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
