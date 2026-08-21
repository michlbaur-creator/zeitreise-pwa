from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[2] / "Dokumente" / "03_Arbeitsstand" / "Suno-Pro_Geraeusche-fuer-Zeitreise.docx"

NAVY = "203748"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B08238"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "F8F1E3"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "66717A"
WHITE = "FFFFFF"
BLACK = "202428"


LOOPS = [
    ("atmo_wind_leicht_loop.wav", "Loop", "gentle natural coastal wind through low vegetation, soft distant air movement, calm and unobtrusive"),
    ("atmo_wind_heiss_loop.wav", "Loop", "dry hot wind over volcanic rock, faint dusty air movement, barren prehistoric landscape"),
    ("atmo_regen_leicht_loop.wav", "Loop", "fine continuous rain over rock and shallow water, delicate narrow raindrops, no thunder"),
    ("atmo_regen_stark_loop.wav", "Loop", "heavy continuous prehistoric rain on rock and ocean water, changing intensity, dense but natural, no thunder"),
    ("atmo_wellen_sanft_loop.wav", "Loop", "very gentle small waves in a warm shallow lagoon, soft water movement, no beach crowd, no birds"),
    ("atmo_brandung_sanft_loop.wav", "Loop", "calm ocean surf washing over a rocky coast, slow recurring waves, soft foam and retreating water"),
    ("atmo_unterwasser_ruhen_loop.wav", "Loop", "quiet underwater ambience in a shallow ancient ocean, muffled water movement, deep calm atmosphere"),
    ("atmo_stroemung_sanft_loop.wav", "Loop", "gentle underwater current flowing past rocks, soft continuous water movement, no animals"),
    ("atmo_vulkanrumpeln_loop.wav", "Loop", "distant low volcanic rumble, subtle earth vibration, irregular and restrained, not explosive"),
    ("atmo_blaetterrauschen_loop.wav", "Loop", "soft leaves rustling in a humid prehistoric forest, gentle irregular breeze"),
    ("atmo_insekten_loop.wav", "Loop", "subtle natural insect ambience in a warm forest, sparse chirps and wing sounds, no modern traffic, no birds"),
    ("atmo_vogelstimmen_loop.wav", "Loop", "peaceful natural woodland bird ambience, only a few small birds, spacious and calm"),
    ("atmo_hummeln_loop.wav", "Loop", "a few bumblebees moving among wild flowers, soft intermittent buzzing, close but not aggressive"),
    ("geraeusch_wasserblasen_loop.wav", "Loop", "small gas bubbles rising through shallow water, irregular delicate bubbling, scientifically natural"),
    ("geraeusch_wasserdampf_zischen_loop.wav", "Loop", "gentle geothermal steam escaping between wet volcanic rocks, soft continuous hiss"),
    ("geraeusch_wasser_plaetschern_loop.wav", "Loop", "small amounts of water softly splashing between rocks, gentle irregular movement"),
]

ONE_SHOTS = [
    ("geraeusch_donner_fern.wav", "One Shot", "single distant thunder roll over an open ocean, soft beginning, low natural rumble, no rain"),
    ("geraeusch_lavablasen.wav", "One Shot", "one thick lava bubble forming and bursting, wet volcanic pop with a short low rumble"),
    ("geraeusch_wassertropfen.wav", "One Shot", "single clear water droplet running over wet stone and falling into a shallow puddle"),
    ("geraeusch_insektenfluegel.wav", "One Shot", "brief delicate flutter of small insect wings passing close to the listener"),
    ("geraeusch_amphibium_quaken_fern.wav", "One Shot", "one distant primitive amphibian call from a quiet swamp, restrained and natural"),
    ("geraeusch_schlamm_schmatzen.wav", "One Shot", "single wet foot movement lifting out of thick shallow mud, soft natural suction"),
    ("geraeusch_krabbeln_stein_leise.wav", "One Shot", "tiny many-legged animal quietly crawling across damp rough stone, delicate close-up foley"),
    ("geraeusch_dinosaurier_schritte_schwer.wav", "One Shot", "one enormous dinosaur footstep on moist forest ground, deep impact, slight earth vibration, no roar"),
    ("geraeusch_dinosaurier_ruf_fern.wav", "One Shot", "single distant large prehistoric animal call, low and natural, documentary realism, not a monster roar"),
    ("geraeusch_rascheln_kleines_tier.wav", "One Shot", "small mammal briefly rustling through dry leaves and low plants, cautious and quiet"),
    ("geraeusch_buntspecht_klopfen.wav", "One Shot", "short sequence of natural woodpecker taps on a tree trunk, distant woodland perspective"),
    ("geraeusch_asteroid_grollen_anschwellend.wav", "One Shot", "deep atmospheric asteroid rumble gradually approaching and increasing for eight seconds, ending before impact"),
    ("einschlag_knall.wav", "One Shot", "one colossal asteroid impact boom, single event only, enormous deep pressure wave, short debris tail, no repeated explosions, no music"),
]


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(element, fill):
    if hasattr(element, "_tc"):
        props = element._tc.get_or_add_tcPr()
    else:
        props = element._p.get_or_add_pPr()
    shd = props.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        props.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run(run, size=11, color=BLACK, bold=False, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tblW = tblPr.find(qn("w:tblW"))
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def set_table_borders(table, color="CAD3DB", size="5"):
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Seite ")
    set_run(run, size=8.5, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_callout(doc, label, text, fill=LIGHT_GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.line_spacing = 1.18
    shade(p, fill)
    r = p.add_run(label + "  ")
    set_run(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run(r, size=10.5, color=BLACK)


def add_prompt_box(doc, label, prompt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.line_spacing = 1.08
    shade(p, LIGHT_GRAY)
    r = p.add_run(label + "\n")
    set_run(r, size=9.5, color=DARK_BLUE, bold=True)
    r = p.add_run(prompt)
    set_run(r, size=9.3, color=BLACK, font="Aptos Mono")


def add_prompt_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ("Dateiname", "Typ", "Text für Suno")
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run(r, size=9.5, color=WHITE, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for row_idx, (filename, sound_type, prompt) in enumerate(rows):
        cells = table.add_row().cells
        if row_idx % 2:
            for cell in cells:
                shade(cell, "F7F9FB")
        values = (filename, sound_type, prompt)
        for idx, value in enumerate(values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run(r, size=8.7 if idx != 2 else 8.9, color=BLACK, bold=(idx == 0), font="Aptos" if idx < 2 else "Aptos Mono")
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_geometry(table, [2232, 1224, 5904])
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 14, MID_GRAY, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1", "Heading 2", "Heading 3") else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if name == "Title":
            p_pr = style._element.get_or_add_pPr()
            border = p_pr.find(qn("w:pBdr"))
            if border is not None:
                p_pr.remove(border)

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("ZEITREISE  |  PRODUKTIONSHILFE")
    set_run(r, size=8.5, color=MID_GRAY, bold=True)
    footer = section.footer
    add_page_number(footer.paragraphs[0])

    # Editorial cover
    for _ in range(4):
        doc.add_paragraph().paragraph_format.space_after = Pt(15)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("PRAXISANLEITUNG")
    set_run(r, size=10, color=GOLD, bold=True)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Geräusche mit Suno Pro erstellen")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("für „Zeitreise - Die Geschichte des Lebens“")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("Atmosphären, Endlosschleifen und Einzelgeräusche\n")
    set_run(r, size=11.5, color=NAVY, bold=True)
    r = p.add_run("Stand: August 2026")
    set_run(r, size=10, color=MID_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("Ruhige Naturdokumentation statt Filmtrailer")
    set_run(r, size=11, color=GOLD, italic=True)

    doc.add_page_break()

    doc.add_heading("1. Schnellstart", level=1)
    intro = doc.add_paragraph()
    intro.add_run("Suno Pro enthält den Bereich ").bold = False
    r = intro.add_run("Sounds")
    r.bold = True
    intro.add_run(" für einzelne Geräusche, Atmosphären und Endlosschleifen. Für die Zeitreise werden keine Songs benötigt.")

    steps = [
        "Suno im Webbrowser öffnen und zu Create wechseln.",
        "Unter Custom den Eintrag Sounds auswählen.",
        "Loop für dauerhafte Geräuschkulissen oder One Shot für einzelne Ereignisse wählen.",
        "Den passenden englischen Text aus dieser Anleitung einfügen.",
        "Mehrere Varianten erzeugen und die ruhigste, natürlichste auswählen.",
        "Die gewählte Variante als WAV herunterladen und mit dem vorgesehenen Dateinamen speichern.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    add_callout(doc, "Wichtig:", "Suno erzeugt üblicherweise zwei Varianten. Nicht die spektakulärste, sondern die glaubwürdigste und zurückhaltendste Fassung auswählen.")

    doc.add_heading("2. Loop oder One Shot?", level=1)
    doc.add_heading("Loop", level=2)
    doc.add_paragraph("Für Regen, Wind, Brandung, Unterwassergeräusche, Waldatmosphäre und andere Klänge, die über längere Zeit unauffällig weiterlaufen sollen.")
    doc.add_heading("One Shot", level=2)
    doc.add_paragraph("Für ein einzelnes Ereignis: Donner, Fußtritt, Knall, Rascheln, Tropfen oder einen kurzen Tierlaut.")

    doc.add_heading("3. Zusätze für zuverlässige Ergebnisse", level=1)
    add_prompt_box(doc, "An jeden Loop anhängen", "natural documentary ambience, seamless loop, 20 seconds, subtle and calm, steady volume, no music, no melody, no voices, no narration, no dramatic cinematic effects")
    add_prompt_box(doc, "An jeden One Shot anhängen", "single isolated sound, no repetition, clean background, no music, no voices, no cinematic score")
    add_prompt_box(doc, "Wenn Suno trotzdem Musik erzeugt", "PURE FIELD RECORDING, SOUND EFFECT ONLY")

    doc.add_page_break()
    doc.add_heading("4. Atmosphären und Endlosschleifen", level=1)
    doc.add_paragraph("Die folgenden Beschreibungen werden als Loop erzeugt. Den allgemeinen Loop-Zusatz jeweils am Ende ergänzen.")
    add_prompt_table(doc, LOOPS)

    doc.add_page_break()
    doc.add_heading("5. Einzelne Geräusche", level=1)
    doc.add_paragraph("Für diese Geräusche One Shot auswählen und den allgemeinen One-Shot-Zusatz anhängen.")
    add_prompt_table(doc, ONE_SHOTS)

    doc.add_page_break()
    doc.add_heading("6. Auswahlregeln für die Zeitreise", level=1)
    checks = [
        "Keine Musik und keine rhythmischen Schläge.",
        "Keine Stimmen oder menschlichen Laute.",
        "Tierstimmen nur entfernt und sparsam einsetzen.",
        "Dinosaurier dürfen nicht wie Filmmonster brüllen.",
        "Endlosschleifen dürfen am Übergang keinen hörbaren Sprung haben.",
        "Regen, Wind und Brandung müssen deutlich leiser als der Sprecher bleiben.",
        "Beim Meteoriten genau einen Einschlag verwenden.",
        "Die zweisekündige Stille in Szene 19 bleibt vollständig still.",
    ]
    for item in checks:
        doc.add_paragraph(" " + item, style="List Bullet")

    doc.add_heading("7. Nachbearbeitung und Ablage", level=1)
    aftercare = [
        "Am Anfang und Ende unnötige Stille entfernen.",
        "Loops mit Kopfhörern mehrfach hintereinander anhören.",
        "Lautstärkesprünge vermeiden; alle Kulissen zunächst eher zu leise einstellen.",
        "Die Dateinamen aus der Medienliste unverändert übernehmen.",
        "Wiederkehrende Atmosphären nur einmal erstellen und in mehreren Szenen verwenden.",
        "Die Suno-Erzeugung und den Download während des aktiven Pro-Abonnements dokumentieren.",
    ]
    for item in aftercare:
        doc.add_paragraph(" " + item, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("8. Rechte und Downloads", level=1)
    doc.add_paragraph("Nach den aktuellen Suno-Hinweisen besitzen zahlende Pro- und Premier-Nutzer kommerzielle Nutzungsrechte an den während des bezahlten Abonnements erzeugten Inhalten. Dies ist keine Garantie dafür, dass ein Werk in jedem Land urheberrechtlich geschützt ist.")
    doc.add_paragraph("WAV-Dateien lassen sich mit Pro oder Premier über die Suno-Webseite herunterladen. Ab dem 3. September 2026 hat Suno für Pro 20 Downloads pro Monat und für Premier 60 Downloads pro Monat angekündigt. Deshalb zuerst die wiederverwendbaren Atmosphären erstellen und nicht für jede Szene eine neue Datei erzeugen.")
    add_callout(doc, "Empfehlung:", "Zu jeder verwendeten Datei Erstellungsdatum, Suno-Link und den endgültigen Dateinamen in einer kleinen Liste festhalten.", fill=LIGHT_BLUE)

    doc.add_heading("9. Offizielle Quellen", level=1)
    sources = [
        "Suno Sounds - Geräusche und Samples: https://help.suno.com/en/articles/10625537",
        "WAV-Downloads: https://help.suno.com/en/articles/2479873",
        "Rechte im bezahlten Tarif: https://help.suno.com/en/articles/9601665",
        "Angekündigte Downloadgrenzen: https://help.suno.com/en/articles/13614785",
    ]
    for source in sources:
        doc.add_paragraph(" " + source, style="List Bullet")

    core = doc.core_properties
    core.title = "Geräusche mit Suno Pro erstellen - Zeitreise"
    core.subject = "Praxisanleitung für Atmosphären, Loops und Einzelgeräusche"
    core.author = "Zeitreise - Die Geschichte des Lebens"
    core.keywords = "Suno, Geräusche, Sounddesign, Zeitreise, PWA"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
