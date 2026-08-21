from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import create_episode2_speaker_draft as base


OUTPUT = Path("Dokumente/03_Arbeitsstand/Zeitreise_Episode2_Kompaktfassung_14_Szenen_V0.1.docx")


SCENES = [
    {
        "id": 1,
        "title": "Der nächste Zeitsprung",
        "time": "Vor etwa 60 bis 55 Millionen Jahren",
        "duration": 42,
        "sources": "bisherige Szenen 1 und 2",
        "beat": "Vom Überleben nach dem Asteroideneinschlag direkt in die Baumkronen; die neue Episode verspricht eine Spurensuche statt einer Aufstiegsleiter.",
        "reuse": "Zeitfelsen, Säugetier-Übergang sowie Baumkronenbilder aus den bisherigen Szenen 1 und 2.",
        "uncertainty": "Keine Detaildebatte über einzelne frühe Primaten im Sprechertext; Vielfalt und Verwandtschaft erscheinen in den Hotspots.",
        "speaker": "Unsere erste Reise endete bei einem kleinen Säugetier. Genau dort geht es weiter. Wälder breiten sich aus, und in den Baumkronen werden Greifhände, gutes räumliches Sehen und ein sicherer Gleichgewichtssinn äußerst praktisch. Wer hier daneben greift, bekommt schließlich keine zweite Probe. Doch halt: Dieses Tier ist nicht einfach ‚der erste Mensch‘. So bequem macht es uns die Evolution nicht. Vor uns liegt kein gerader Weg, sondern ein gewaltiger Stammbaum – mit vielen Ästen, überraschenden Abzweigungen und einigen Verwandten, von denen nur noch Zähne und Knochen übrig sind. Willkommen bei der Spurensuche nach uns selbst.",
    },
    {
        "id": 2,
        "title": "Die Welt der Primaten",
        "time": "Vor etwa 55 bis 10 Millionen Jahren",
        "duration": 40,
        "sources": "bisherige Szenen 2 und 3",
        "beat": "Aus den Anpassungen an das Baumleben entsteht eine enorme Primaten- und Menschenaffenvielfalt – ohne bevorzugten Hauptzweig.",
        "reuse": "Baumkronen, Primatengruppe und Miozän-Landschaft aus den bisherigen Szenen 2 und 3.",
        "uncertainty": "Einzelne fossile Menschenaffen und ihre umstrittenen Verwandtschaftsverhältnisse bleiben den Hotspots vorbehalten.",
        "speaker": "In den Bäumen zählt jeder Griff. Viele frühe Primaten besitzen deshalb bewegliche Hände und Füße, empfindliche Fingerspitzen und nach vorn gerichtete Augen. Eine Art eingebautes Abstandsmessgerät – nur ohne Batterieanzeige. Millionen Jahre später ist aus dieser Verwandtschaft eine erstaunliche Vielfalt geworden. In Afrika, Asien und Europa leben zahlreiche Menschenaffen: manche klettern, manche bewegen sich häufiger am Boden, viele verschwinden wieder. Keiner trägt ein Schild mit der Aufschrift ‚Hier entlang zum Menschen‘. Unsere Linie ist nur ein Zweig unter vielen. Evolution verteilt keine Pokale. Sie probiert aus – und behält, was gerade funktioniert.",
    },
    {
        "id": 3,
        "title": "Getrennte Wege",
        "time": "Vor ungefähr 10 bis 6 Millionen Jahren",
        "duration": 42,
        "sources": "bisherige Szenen 4 und 5",
        "beat": "Wechselnde afrikanische Lebensräume rahmen die Trennung der Linien von Menschen sowie Schimpansen und Bonobos.",
        "reuse": "Landschaftsmosaik und verzweigter Stammbaum aus den bisherigen Szenen 4 und 5.",
        "uncertainty": "Der genaue gemeinsame Vorfahr wird nicht rekonstruiert; im Hotspot steht ausdrücklich, dass sein Aussehen unbekannt ist.",
        "speaker": "Afrika wird jetzt zum Mosaik aus Wald, Buschland, Flussufern und offenen Flächen. Lebensräume verändern sich – und ihre Bewohner verändern sich mit. In dieser Zeit trennen sich auch die Wege unserer Vorfahren und der Vorfahren heutiger Schimpansen und Bonobos. Wichtig: Wir stammen nicht von den heutigen Schimpansen ab. Wir teilen mit ihnen einen älteren gemeinsamen Vorfahren. Wie der aussah, wissen wir nicht; Familienfotos waren damals noch erstaunlich selten. Von nun an entwickeln sich mehrere Linien gleichzeitig weiter. Keine führt von ‚primitiv‘ zu ‚fortschrittlich‘. Es sind einfach verschiedene Familiengeschichten – und unsere beginnt gerade erst interessant zu werden.",
    },
    {
        "id": 4,
        "title": "Der rätselhafte Gang auf zwei Beinen",
        "time": "Vor etwa 7 bis 4,4 Millionen Jahren",
        "duration": 42,
        "sources": "bisherige Szenen 6 und 7",
        "beat": "Sahelanthropus und Ardi zeigen gemeinsam: Zweibeinigkeit entstand schrittweise und lange neben dem Klettern.",
        "reuse": "Sahelanthropus-Landschaft, Ardi-Bild und vorhandene Skelett- beziehungsweise Gang-Overlays.",
        "uncertainty": "Die genaue Gangart und Stammbaumposition bleiben als Forschungsfrage sichtbar, werden aber nur einmal knapp benannt.",
        "speaker": "Die nächste Spur steckt in ein paar sehr alten Knochen. Bei Sahelanthropus sprechen neue Untersuchungen dafür, dass er sich vor rund sieben Millionen Jahren zumindest zeitweise auf zwei Beinen bewegte. Einige Millionen Jahre später zeigt Ardi eine verblüffende Mischung: Ihr Becken passt zum aufrechten Gehen, ihr großer Zeh kann noch kräftig zugreifen. Praktisch, wenn der Boden interessant ist, der nächste Baum aber sicherer wirkt. Die Aktenlage bleibt dünn, doch eine Botschaft ist klar: Zweibeinigkeit wurde nicht an einem Dienstag fertig erfunden. Sie entstand Schritt für Schritt – während das Klettern noch lange zum Programm gehörte.",
    },
    {
        "id": 5,
        "title": "Lucy und die Spuren in der Asche",
        "time": "Vor etwa 3,6 bis 3,2 Millionen Jahren",
        "duration": 44,
        "sources": "bisherige Szenen 8 und 9",
        "beat": "Fußspuren und Skelett werden zu zwei sich ergänzenden Zeugen für aufrechtes Gehen und fortbestehendes Klettern.",
        "reuse": "Laetoli-Spuren und Lucy-Bild; beide Motive können in einer ruhigen Überblendung verbunden werden.",
        "uncertainty": "Die Zuordnung der Laetoli-Spuren zu Australopithecus afarensis bleibt als wahrscheinliche, nicht absolute Aussage formuliert.",
        "speaker": "Jetzt melden sich zwei Zeugen. Der erste hinterlässt vor 3,6 Millionen Jahren Fußspuren in feuchter Vulkanasche von Laetoli. Ferse, Fußgewölbe, großer Zeh: Hier geht jemand aufrecht. Wer genau, bleibt offen – wahrscheinlich Australopithecus afarensis. Der zweite Zeuge heißt Lucy. Ihr Skelett zeigt ebenfalls regelmäßiges Gehen auf zwei Beinen. Lange Arme und gebogene Finger verraten aber: Klettern war weiterhin nützlich. Und ihr Gehirn ist deutlich kleiner als unseres. Erst kamen also die aufrechten Schritte, viel später das große Gehirn. Die Evolution hält sich offenbar nicht an die Reihenfolge, die wir für logisch halten.",
    },
    {
        "id": 6,
        "title": "Der Stein, der alles durcheinanderbringt",
        "time": "Früheste bekannte Funde vor etwa 3,3 Millionen Jahren",
        "duration": 45,
        "sources": "bisherige Szene 10",
        "beat": "Erster großer Erkenntnismoment: Ein Werkzeugfund ist älter als die frühesten sicheren Homo-Fossilien und entlarvt die einfache Fortschrittsgeschichte.",
        "reuse": "Vollständiges Werkzeugbild und Schlaganimation der bisherigen Szene 10.",
        "uncertainty": "Der unbekannte Hersteller ist die Pointe des Quiz, nicht eine wiederholte Bremse im Text.",
        "speaker": "Ein Schlag – und plötzlich stimmt die vertraute Geschichte nicht mehr. In Lomekwi liegen bearbeitete Steine, rund 3,3 Millionen Jahre alt. Das Problem: Sie sind älter als die frühesten sicher bekannten Fossilien der Gattung Homo. Wer hat sie hergestellt? Vielleicht eine Australopithecus-Art, vielleicht eine andere Homininenform. Der Erfinder hat seinen Namen leider nicht in den Stein geritzt. Sicher ist nur: Jemand wählte Material aus, hielt es fest und veränderte es mit wiederholten Schlägen. Das wirkt grob, verlangt aber Planung und Übung. Werkzeugtechnik begann offenbar, bevor unsere eigene Gattung überhaupt eindeutig auf der Bühne stand.",
    },
    {
        "id": 7,
        "title": "Homo kommt in Bewegung",
        "time": "Seit etwa 2,8 Millionen Jahren",
        "duration": 42,
        "sources": "bisherige Szenen 11 und 12",
        "beat": "Die unscharfen Anfänge der Gattung Homo führen direkt zu Homo erectus und einem Körper für weite Wege.",
        "reuse": "Fossil- und Gruppenbild der bisherigen Szene 11 sowie Homo-erectus-Landschaft aus Szene 12.",
        "uncertainty": "Taxonomische Grenzfragen werden in einen Hotspot ausgelagert; im Haupttext zählt das Mosaik der Veränderungen.",
        "speaker": "Mit der Gattung Homo beginnt kein sauber beschriftetes neues Kapitel. Der älteste Kiefer, den Forschende zu Homo stellen, ist etwa 2,8 Millionen Jahre alt. Doch Zähne, Hände, Gehirn und Körper verändern sich nicht im Gleichschritt – die Evolution führt offenbar keine ordentliche Checkliste. Später erscheint Homo erectus: längere Beine, kürzere Arme, ein Körper für weite Strecken. Sein Gehirn ist größer als bei früheren Homininen, aber entscheidend ist das Zusammenspiel aus Körper, Werkzeugen, Nahrung, Lernen und Gruppe. Kein einzelnes Merkmal macht plötzlich den Menschen. Doch zusammen bringen diese Veränderungen unsere Verwandten weit voran – im wahrsten Sinne des Wortes.",
    },
    {
        "id": 8,
        "title": "Die erste große Reise",
        "time": "Vor etwa 1,85 bis 1,7 Millionen Jahren",
        "duration": 42,
        "sources": "bisherige Szene 13",
        "beat": "Großer Mittelübergang: Frühe Menschen verlassen Afrika nicht als Expedition, sondern generationenweise entlang günstiger Lebensräume.",
        "reuse": "Das vorhandene Wanderungsbild aus Szene 13; anschließend ruhige Kartenbewegung als Mittelübergang.",
        "uncertainty": "Einzelne Routen bleiben bewusst offen; die belastbare zeitliche Aussage steht im Vordergrund.",
        "speaker": "Vor ungefähr 1,8 Millionen Jahren tauchen frühe Menschen außerhalb Afrikas auf. In Dmanisi, im heutigen Georgien, leben kleine Gruppen zwischen Grasland, Wald und Bergen. Später reichen Spuren bis weit nach Ostasien. Das war wahrscheinlich keine große Expedition mit Reiseleiter und festem Ziel. Generation für Generation folgten Menschen Wasser, Nahrung und günstigen Lebensräumen. Manche Wege endeten, andere führten weiter. Auf unserer Karte sieht das nach einer gewaltigen Wanderung aus. Für jeden Einzelnen war es vermutlich nur der nächste erreichbare Ort – und die uralte Frage: Gibt es dort vorne vielleicht etwas zu essen?",
    },
    {
        "id": 9,
        "title": "Feuer verändert den Alltag",
        "time": "Sicher kontrollierte Feuerstellen spätestens vor etwa 790.000 Jahren",
        "duration": 42,
        "sources": "bisherige Szene 14",
        "beat": "Feuer wird als Bündel mehrerer Veränderungen erzählt: Nahrung, Schutz, Licht und sozialer Treffpunkt.",
        "reuse": "Feuerbild und vorhandene Flammen-, Rauch- und Lichtbewegungen der bisherigen Szene 14.",
        "uncertainty": "Ältere Brandspuren werden im Hotspot erläutert; der Sprechertext trennt klar zwischen Feuer und kontrollierter Feuerstelle.",
        "speaker": "Feuer ist Wärme, Licht und Schutz. Es macht manche Nahrung leichter verdaulich und schafft einen Ort, an dem eine Gruppe zusammenkommt. Vermutlich ist es auch der erste Treffpunkt, an dem alle behaupten, jemand anderes solle Holz nachlegen. Sicher kontrollierte Feuerstellen kennen wir spätestens vor etwa 790.000 Jahren. Es gibt ältere Brandspuren, doch ein schwarzer Fleck ist noch keine Küche – dort könnte auch ein Busch gebrannt haben. Entscheidend ist die wiederholte, kontrollierte Nutzung. Erst dann wird aus einer Flamme ein Werkzeug, das Nahrung, Sicherheit und wahrscheinlich auch das gemeinsame Lernen verändert.",
    },
    {
        "id": 10,
        "title": "Eine Welt voller Menschenformen",
        "time": "Vor etwa 700.000 bis 300.000 Jahren",
        "duration": 44,
        "sources": "bisherige Szene 15",
        "beat": "Ein großes Panorama ersetzt mehrere taxonomische Zwischenstationen und bereitet Neandertaler, Denisova und Homo sapiens vor.",
        "reuse": "Panoramabild und verzweigte Linien der bisherigen Szene 15.",
        "uncertainty": "Strittige Artnamen werden nicht aufgezählt; ein Hotspot erklärt, warum Fossilien nicht immer in klare Schubladen passen.",
        "speaker": "Je näher wir unserer Gegenwart kommen, desto voller wird die Welt der Menschenformen. Populationen leben in Afrika, Europa und Asien. Sie trennen sich, verändern sich und begegnen sich später wieder. Früher ordnete man viele Fossilien gern in eine saubere Reihe – schön übersichtlich, nur wahrscheinlich falsch. Heute sieht der Stammbaum eher wie ein Gebüsch aus. Manche Funde passen zu keiner Art ganz eindeutig, und selbst Fachleute beschriften dieselbe Schublade unterschiedlich. Das liegt nicht daran, dass die Evolution unordentlich gearbeitet hätte. Die Schubladen stammen von uns. Sicher ist: Mehrere Menschenformen leben lange gleichzeitig. Unsere Bühne wird jetzt ziemlich voll.",
    },
    {
        "id": 11,
        "title": "Die Neandertaler",
        "time": "Vor etwa 400.000 bis 40.000 Jahren",
        "duration": 46,
        "sources": "bisherige Szene 16",
        "beat": "Das Klischee vom groben Höhlenmenschen wird durch Können, Anpassung und Fürsorge ersetzt.",
        "reuse": "Neandertaler-Lager und Tätigkeiten der bisherigen Szene 16.",
        "uncertainty": "Rituale und symbolische Handlungen bleiben aus dem Haupttext heraus; gut belegte Tätigkeiten tragen die Szene.",
        "speaker": "Neandertaler sind keine unbeholfenen Vorstufen von uns – und keine Statisten, die nur auf unseren Auftritt warten. Hunderttausende Jahre leben sie in Europa und Westasien, stellen anspruchsvolle Werkzeuge her, nutzen Feuer und kümmern sich offenbar um verletzte oder kranke Gruppenmitglieder. Ihr kräftiger Körper passt gut zu kalten Bedingungen. Wer im eiszeitlichen Europa lebt, kann schließlich nicht einfach die Heizung höher drehen. Ihre Kulturen verändern sich über Zeit und Raum. Neandertaler sind ein eigener, erfolgreicher Menschenzweig. Und völlig getrennt von uns waren sie ebenfalls nicht. Davon tragen viele Menschen heute noch kleine Spuren im eigenen Erbgut.",
    },
    {
        "id": 12,
        "title": "Denisova – eine Menschenlinie aus einer Schachtel",
        "time": "Vor mehreren hunderttausend bis vor mindestens etwa 50.000 Jahren",
        "duration": 48,
        "sources": "bisherige Szene 17",
        "beat": "Die Spurensuche erreicht ihren stärksten Detektivmoment: DNA und Proteine machen aus wenigen Resten eine weit verbreitete Menschenlinie.",
        "reuse": "Fingerknochen, Zahn, Harbin-Schädel und Molekül-Overlays der bisherigen Szene 17.",
        "uncertainty": "Regionale Vielfalt und Aussehen bleiben offen; der Harbin-Schädel erscheint als sehr wahrscheinlicher, nicht endgültiger Gesamtschlüssel.",
        "speaker": "Manchmal passt eine ganze Menschenlinie in eine winzige Schachtel. Denisova-Menschen wurden zuerst durch DNA aus einem kleinen Fingerknochen erkannt – weniger Material, als man für ein ordentliches Puzzle erwarten würde. Lange kannten wir fast nur Zähne und Bruchstücke. Seit 2025 liefern alte Proteine und Erbgut aus Zahnstein einen größeren Hinweis: Auch der mehr als 146.000 Jahre alte Harbin-Schädel gehört sehr wahrscheinlich zu dieser Linie. Das Bild wird klarer, bleibt aber lückenhaft. Genetische Spuren zeigen, dass Denisova-Populationen weit in Asien lebten und anderen Menschenformen begegneten. Hier erzählt die DNA eindeutig mehr als die Knochenkiste.",
    },
    {
        "id": 13,
        "title": "Homo sapiens trifft Verwandte",
        "time": "Seit mindestens etwa 300.000 Jahren",
        "duration": 60,
        "sources": "bisherige Szenen 18 und 19",
        "beat": "Zweiter großer Erkenntnismoment: Unsere Art entsteht in einem afrikanischen Netzwerk und trifft später auf bereits lebende Menschenformen.",
        "reuse": "Afrika-Landschaft aus Szene 18, Wasserstellen-Begegnung aus Szene 19 sowie vorhandene DNA- und Verbindungslinien.",
        "uncertainty": "Konkrete Begegnungsabläufe bleiben unerzählt; Sprechertext trennt gesicherte genetische Vermischung von unbekannten persönlichen Situationen.",
        "speaker": "Auch unsere eigene Art erscheint nicht plötzlich an einem einzigen Geburtsort. Rund 300.000 Jahre alte Fossilien aus Marokko und weitere Funde aus Afrika zeigen ein Netzwerk verschiedener Populationen. Homo sapiens entsteht Schritt für Schritt – eher aus vielen verbundenen Punkten als aus einer einzigen Wiege. Später breiten sich Gruppen weiter aus. Doch die Welt ist nicht leer: Sie treffen auf Neandertaler und Denisova-Populationen. Manche gehen einander vielleicht aus dem Weg, manche konkurrieren, manche bekommen gemeinsame Kinder. Was genau geschah, verrät kein Fossil. Das Erbgut heutiger Menschen beweist jedoch: Einige dieser Linien haben sich vermischt. Unsere Familiengeschichte ist also nicht nur verzweigt. An manchen Stellen wachsen die Zweige wieder zusammen.",
    },
    {
        "id": 14,
        "title": "Eine Menschheit",
        "time": "Heute",
        "duration": 50,
        "sources": "bisherige Szene 20",
        "beat": "Ruhiges Finale am Zeitfelsen: gemeinsame Herkunft, reale Vielfalt und keine biologische Rangordnung.",
        "reuse": "Das neue Finale mit Kinderhand und Zeitfelsen bleibt vollständig erhalten.",
        "uncertainty": "Keine zusätzliche Forschungsdebatte im Finale; die Aussage zu menschlicher Variation folgt dem wissenschaftlichen Konsens.",
        "speaker": "Unsere Geschichte ist verzweigt, voller Seitenwege und Begegnungen. Viele Menschenformen leben lange gleichzeitig. Einige verschwinden, andere hinterlassen Spuren in unserem Erbgut. Heute gibt es nur noch eine Menschenart: Homo sapiens. Unsere sichtbare Vielfalt ist real und wertvoll, doch sie teilt uns nicht in biologische Menschenrassen. Dafür sind wir viel zu eng miteinander verwandt – selbst wenn Familienfeiern gelegentlich einen anderen Eindruck vermitteln. Letztlich stammen wir alle aus afrikanischen Populationen. Deine Hand liegt wieder auf dem Zeitfelsen. Die Spurensuche endet hier – vorerst. Denn was Menschen später mit ihrer Welt machen, ist ein neues Kapitel. Und wie du dir denken kannst, wird auch dieses nicht gerade kurz.",
    },
]


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run(kicker.add_run("ZEITREISE – DIE GESCHICHTE DES LEBENS"), size=10, color=base.GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run(title.add_run("Episode 2"), size=30, color=base.NAVY, bold=True)
    title.paragraph_format.space_after = Pt(5)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run(subtitle.add_run("Die Entwicklung des Menschen"), size=19, color=base.BLUE)
    subtitle.paragraph_format.space_after = Pt(3)
    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run(strap.add_run("Eine verzweigte Geschichte"), size=14, color=base.MUTED, italic=True)
    strap.paragraph_format.space_after = Pt(28)
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run(version.add_run("Kompaktfassung · 14 Szenen · V0.1"), size=11, color=base.GOLD, bold=True)
    version.paragraph_format.space_after = Pt(88)
    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_run(status.add_run("Alternative Prüffassung zur 20-Szenen-Version\nNoch nicht in die App übernommen"), size=10, color=base.MUTED, italic=True)
    doc.add_page_break()


def add_concept(doc):
    doc.add_heading("Das neue Erzählprinzip", level=1)
    lead = doc.add_paragraph()
    base.set_run(lead.add_run("Nicht weniger Wissenschaft – sondern weniger Wiederholung."), size=13, color=base.NAVY, bold=True)
    lead.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.add_run("Die Episode wird als Spurensuche erzählt. ").bold = True
    p.add_run("Jede Szene beginnt mit einem Fund oder einer Beobachtung, führt zu einer überraschenden Erkenntnis und öffnet die nächste Frage. Unsicherheiten bleiben ehrlich, unterbrechen aber nicht mehr in jeder Szene den Erzählfluss.")
    principles = [
        "14 Szenen statt 20; voraussichtliche reine Sprecherzeit rund zehn Minuten.",
        "Durchgängige neugierige, trockene Erzählstimme statt sachlicher Absätze mit nachträglich eingesetzten Pointen.",
        "Forschungsfragen werden nur dort gesprochen, wo sie die Spannung tragen; Details wandern in Hotspots.",
        "Zwei verbindliche Quizstopps: Szene 6 beim unbekannten Werkzeughersteller und Szene 13 bei der genetischen Vermischung.",
        "Musik weiterhin sparsam: Beginn, großer Übergang vor beziehungsweise in Szene 8 und Finale.",
        "Landwirtschaft bleibt vollständig einer späteren Episode vorbehalten.",
    ]
    for item in principles:
        doc.add_paragraph(item, style="List Bullet")
    total = sum(scene["duration"] for scene in SCENES)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    base.set_run(p.add_run("Geplante Sprecherzeit: "), size=10.5, color=base.GOLD, bold=True)
    base.set_run(p.add_run(f"etwa {total // 60} Minuten {total % 60} Sekunden; Übergänge und Quiz kommen hinzu."), size=10.5, color=base.MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    base.set_run(p.add_run("Sicherheitsnetz: "), size=10.5, color=base.GOLD, bold=True)
    base.set_run(p.add_run("Die vorhandene 20-Szenen-Fassung, alle Bilder und die App bleiben unangetastet, bis diese Alternative freigegeben ist."), size=10.5, color=base.MUTED)
    doc.add_page_break()


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    base.set_run(p.add_run(f"{label}: "), size=9.5, color=base.GOLD, bold=True)
    base.set_run(p.add_run(text), size=9.5, color=base.MUTED)
    return p


def add_scene(doc, scene):
    heading = doc.add_heading(f"{scene['id']}. {scene['title']}", level=1)
    heading.paragraph_format.keep_with_next = True
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(8)
    meta.paragraph_format.keep_with_next = True
    base.set_run(meta.add_run(scene["time"]), size=9.5, color=base.GOLD, bold=True)
    words = len(scene["speaker"].split())
    base.set_run(meta.add_run(f"   ·   ca. {scene['duration']} Sekunden   ·   {words} Wörter   ·   {scene['sources']}"), size=9.5, color=base.MUTED)
    add_labeled_paragraph(doc, "Dramaturgischer Zug", scene["beat"])

    table = doc.add_table(rows=1, cols=1)
    base.set_repeat_table_width(table)
    cell = table.cell(0, 0)
    base.set_cell_shading(cell, base.PALE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    base.set_run(paragraph.add_run(scene["speaker"]), size=11.2, color=base.NAVY)

    reuse = add_labeled_paragraph(doc, "Vorhandene Bilder", scene["reuse"])
    reuse.paragraph_format.space_before = Pt(8)
    add_labeled_paragraph(doc, "Wissenschaftliche Vertiefung", scene["uncertainty"])
    doc.add_page_break()


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    base.configure_document(doc)
    header = doc.sections[0].header.paragraphs[0]
    header.text = "ZEITREISE  ·  EPISODE 2  ·  KOMPAKTFASSUNG"
    base.set_run(header.runs[0], size=8.5, color=base.MUTED, bold=True)
    add_cover(doc)
    add_concept(doc)
    for index, scene in enumerate(SCENES):
        add_scene(doc, scene)
    # Remove the final empty page-break paragraph.
    last = doc.paragraphs[-1]
    if last._p.getparent() is not None:
        last._p.getparent().remove(last._p)
    props = doc.core_properties
    props.title = "Zeitreise – Episode 2 – Kompaktfassung mit 14 Szenen V0.1"
    props.subject = "Alternative dramaturgische Prüffassung für Episode 2"
    props.author = "Michael Baur / Codex"
    props.keywords = "Zeitreise, Evolution, Menschheitsentwicklung, Kompaktfassung, Sprechertexte"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
